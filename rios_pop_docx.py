"""
RIOS — Gerador de POP em Word (.docx)
Reescrito em 01/09/2026 para reproduzir EXATAMENTE o padrão visual dos POPs de
referência entregues pelo Chef (POP_COZ_001 e POP_COZ_002, Mestre do Pão):
estilos nomeados do Word (POP Heading / POP Body / POP Bullet), cabeçalho
institucional único (repetido em toda página, sem "capa" separada), tabelas
padronizadas para Responsabilidades / Materiais / PCC / Indicadores, rodapé
de documento controlado, e fluxograma nativo do Word (caixas + setas "↓"),
sem imagem PNG.
"""
import os
import re
import unicodedata

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

DEFAULT_COLOR = "5E4778"
FONT = "Arial"


# ── Helpers de cor / texto ──────────────────────────────────────────────
def norm_color(hex_color):
    h = (hex_color or DEFAULT_COLOR).lstrip('#').strip()
    if len(h) != 6 or any(c not in '0123456789abcdefABCDEF' for c in h):
        h = DEFAULT_COLOR
    return h.upper()


def rgb_of(hex_color):
    h = norm_color(hex_color)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def tint_of(hex_color, amount=0.10):
    """Mistura a cor institucional com branco (~10%) — usado nas caixas
    intermediárias do fluxograma, igual ao padrão de referência (5E4778 -> EEEAF2)."""
    h = norm_color(hex_color)
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    r2 = int(r + (255 - r) * (1 - amount))
    g2 = int(g + (255 - g) * (1 - amount))
    b2 = int(b + (255 - b) * (1 - amount))
    return f"{r2:02X}{g2:02X}{b2:02X}"


def slugify(text, maxlen=40):
    text = str(text or "POP")
    text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii')
    text = re.sub(r'[^A-Za-z0-9]+', '_', text).strip('_')
    return (text or "POP")[:maxlen]


# ── Helpers OOXML de baixo nível ────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), norm_color(hex_color))
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_borders(cell, sz=6, color="777777", edges=('top', 'left', 'bottom', 'right')):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
    for edge in edges:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)


def set_cell_margins(cell, top=90, bottom=90, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        el = OxmlElement(f'w:{tag}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def set_cell_vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), 'center')
    tcPr.append(va)


def no_table_autofit(table):
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)


def set_col_widths(table, widths_cm):
    no_table_autofit(table)
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)
    grid = table._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        for idx, gridcol in enumerate(grid.findall(qn('w:gridCol'))):
            if idx < len(widths_cm):
                gridcol.set(qn('w:w'), str(int(widths_cm[idx] * 567)))


def add_field(paragraph, field_code):
    """Insere um campo de Word (ex.: PAGE, NUMPAGES) — atualiza sozinho no Word."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' {field_code} '
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    txt = OxmlElement('w:t')
    txt.text = '1'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run2 = paragraph.add_run(); run2._r.append(instr)
    run3 = paragraph.add_run(); run3._r.append(fld_sep)
    run3._r.append(txt)
    run4 = paragraph.add_run(); run4._r.append(fld_end)


def set_keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    el = OxmlElement('w:keepNext')
    pPr.append(el)


def make_placeholder_logo(out_path, text="LOGOTIPO\nDO CLIENTE", size=(600, 300)):
    img = Image.new("RGB", size, "white")
    from PIL import ImageDraw, ImageFont
    d = ImageDraw.Draw(img)
    d.rectangle([4, 4, size[0]-4, size[1]-4], outline=(150, 150, 150), width=3)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 30)
    except Exception:
        font = ImageFont.load_default()
    lines = text.split("\n")
    line_h = getattr(font, 'size', 20) + 8
    total_h = line_h * len(lines)
    y = (size[1]-total_h)//2
    for line in lines:
        bbox = d.textbbox((0, 0), line, font=font)
        w = bbox[2]-bbox[0]
        d.text(((size[0]-w)//2, y), line, font=font, fill=(150, 150, 150))
        y += line_h
    img.save(out_path)
    return out_path


# ── Estilos nomeados do Word (POP Heading / POP Body / POP Bullet) ──────
def register_pop_styles(doc, color_hex):
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(9)

    heading = styles.add_style('POP Heading', WD_STYLE_TYPE.PARAGRAPH)
    heading.base_style = normal
    heading.font.name = FONT
    heading.font.bold = True
    heading.font.size = Pt(12)
    heading.font.color.rgb = rgb_of(color_hex)
    heading.paragraph_format.space_before = Pt(9)
    heading.paragraph_format.space_after = Pt(4)

    body = styles.add_style('POP Body', WD_STYLE_TYPE.PARAGRAPH)
    body.base_style = normal
    body.font.name = FONT
    body.font.size = Pt(9)
    body.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    body.paragraph_format.space_after = Pt(5)

    bullet = styles.add_style('POP Bullet', WD_STYLE_TYPE.PARAGRAPH)
    bullet.base_style = normal
    bullet.font.name = FONT
    bullet.font.size = Pt(9)
    bullet.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.left_indent = Cm(0.45)
    bullet.paragraph_format.first_line_indent = Cm(-0.25)

    return heading, body, bullet


DESTAQUE_TAGS = {
    "critico": (" [PONTO CRÍTICO]", RGBColor(0xB0, 0x00, 0x20)),
    "regra": (" [REGRA PRINCIPAL]", RGBColor(0xC9, 0x74, 0x0A)),
    "financeiro": (" [PCC FINANCEIRO]", RGBColor(0xA6, 0x8B, 0x00)),
}


# ── Construção do documento ─────────────────────────────────────────────
def build_pop_docx(pop, cliente, out_path, logo_path=None, color_hex=None, tmp_dir=None):
    tmp_dir = tmp_dir or os.path.dirname(out_path) or "."
    color_hex = norm_color(color_hex or (cliente or {}).get('corHex') or DEFAULT_COLOR)
    cliente_nome = (cliente or {}).get('nome') or pop.get('cliente') or 'Cliente'
    sistema_nome = (cliente or {}).get('sistemaGestao') or f"Sistema de Gestão {cliente_nome}"

    revisoes = pop.get('revisoes') or []
    ultima_rev = revisoes[-1] if revisoes else {}
    rev_num = str(ultima_rev.get('numero', 0)).zfill(2)
    rev_data = ultima_rev.get('data') or pop.get('atualizadoEm', '')[:10] or ''

    codigo = pop.get('codigo') or 'POP-0000'
    titulo = pop.get('titulo') or 'Procedimento Operacional Padrão'
    setor = pop.get('setor') or '—'

    # logo real ou placeholder
    logo_final = logo_path
    if not logo_final or not os.path.exists(logo_final):
        logo_final = os.path.join(tmp_dir, "_placeholder_logo.png")
        make_placeholder_logo(logo_final)

    doc = Document()
    register_pop_styles(doc, color_hex)

    # metadados do documento (título/assunto/autor) — igual ao padrão de referência
    cp = doc.core_properties
    cp.title = f"{codigo} - {titulo}"
    cp.subject = f"Procedimento Operacional Padronizado - {cliente_nome}"
    cp.author = "Chef Marco À Souza"

    # ── Seção única — página A4, margens e cabeçalho institucional fixo ──
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)
    sec.top_margin = Cm(3.75)
    sec.bottom_margin = Cm(1.5)
    sec.header_distance = Cm(1.27)
    sec.footer_distance = Cm(1.27)
    usable_cm = 21.0 - 1.8 - 1.8

    # cabeçalho: tabela de 3 colunas (logo | sistema+título | código/rev/data/páginas)
    header = sec.header
    for hp in header.paragraphs:
        hp.text = ""
    htbl = header.add_table(rows=1, cols=3, width=Cm(usable_cm))
    htbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    w_logo = usable_cm * 0.18
    w_mid = usable_cm * 0.60
    w_meta = usable_cm - w_logo - w_mid
    set_col_widths(htbl, [w_logo, w_mid, w_meta])
    hc_logo, hc_mid, hc_meta = htbl.rows[0].cells
    for c in (hc_logo, hc_mid, hc_meta):
        set_cell_borders(c, sz=4, color="000000")
        set_cell_margins(c, top=60, bottom=60, left=80, right=80)
        set_cell_vcenter(c)

    p_logo = hc_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run()
    with Image.open(logo_final) as im:
        iw, ih = im.size
    max_w_cm, max_h_cm = 2.4, 2.2
    disp_w = max_w_cm
    disp_h = max_w_cm * ih / iw if iw else max_h_cm
    if disp_h > max_h_cm:
        disp_h = max_h_cm
        disp_w = max_h_cm * iw / ih if ih else max_w_cm
    r_logo.add_picture(logo_final, width=Cm(disp_w), height=Cm(disp_h))

    p_sis = hc_mid.paragraphs[0]
    p_sis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sis = p_sis.add_run(sistema_nome)
    r_sis.bold = True; r_sis.font.size = Pt(11); r_sis.font.name = FONT
    p_tit = hc_mid.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tit = p_tit.add_run(titulo.upper())
    r_tit.font.size = Pt(11); r_tit.font.name = FONT
    r_tit.font.color.rgb = rgb_of(color_hex)

    meta_lines = [("Cód.: ", codigo), ("Rev.: ", rev_num), ("Data: ", rev_data or '—')]
    first = True
    for label, value in meta_lines:
        pm = hc_meta.paragraphs[0] if first else hc_meta.add_paragraph()
        first = False
        rl = pm.add_run(label); rl.bold = True; rl.font.size = Pt(8); rl.font.name = FONT
        rv = pm.add_run(value); rv.font.size = Pt(8); rv.font.name = FONT
    pm = hc_meta.add_paragraph()
    rl = pm.add_run("Páginas: "); rl.bold = True; rl.font.size = Pt(8); rl.font.name = FONT
    rv = pm.add_run("1 / "); rv.font.size = Pt(8); rv.font.name = FONT
    add_field(pm, "NUMPAGES")
    for r in pm.runs[-3:]:
        r.font.size = Pt(8); r.font.name = FONT

    # rodapé: linha de "documento controlado"
    footer = sec.footer
    for fp in footer.paragraphs:
        fp.text = ""
    p_foot = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run(f"Documento controlado • {cliente_nome} • {codigo} • Rev. {rev_num}")
    r_foot.font.size = Pt(7); r_foot.font.name = FONT
    r_foot.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Página 1 — Histórico de Revisão ─────────────────────────────
    ph = doc.add_paragraph(style='POP Heading')
    ph.paragraph_format.space_before = Pt(16)
    ph.add_run("Histórico de Revisão")

    rev_tbl = doc.add_table(rows=1, cols=5)
    set_col_widths(rev_tbl, [1.8, usable_cm - 1.8 - 4.2 - 4.2 - 3.0, 4.2, 4.2, 3.0])
    hdr = rev_tbl.rows[0].cells
    headers0 = ["Revisão", "Descrição", "Revisado por", "Aprovado por", "Data"]
    for i, htext in enumerate(headers0):
        set_cell_bg(hdr[i], color_hex)
        set_cell_borders(hdr[i], sz=6, color="777777")
        set_cell_margins(hdr[i])
        set_cell_vcenter(hdr[i])
        pp = hdr[i].paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pp.add_run(htext); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
        rr.font.size = Pt(8); rr.font.name = FONT

    rows_data = revisoes if revisoes else [{"numero": 0, "descricao": "Revisão inicial",
                                             "revisadoPor": "", "aprovadoPor": "", "data": ""}]
    for rdata in rows_data:
        row = rev_tbl.add_row().cells
        vals = [str(rdata.get('numero', 0)).zfill(2), rdata.get('descricao', ''),
                rdata.get('revisadoPor', ''), rdata.get('aprovadoPor', ''), rdata.get('data', '')]
        for i, v in enumerate(vals):
            set_cell_borders(row[i], sz=6, color="777777")
            set_cell_margins(row[i])
            set_cell_vcenter(row[i])
            pp = row[i].paragraphs[0]
            rr = pp.add_run(v); rr.font.size = Pt(8); rr.font.name = FONT

    p_gap = doc.add_paragraph()
    p_gap.paragraph_format.space_before = Pt(64)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("PROCEDIMENTO OPERACIONAL PADRONIZADO")
    r_sub.font.size = Pt(11); r_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66); r_sub.font.name = FONT

    doc.add_page_break()

    # ── Corpo: seções numeradas ──────────────────────────────────────
    counter = {"n": 0}

    def add_heading(title):
        counter["n"] += 1
        p = doc.add_paragraph(style='POP Heading')
        set_keep_with_next(p)
        p.add_run(f"{counter['n']}. {title}")
        return p

    def add_body_text(text):
        p = doc.add_paragraph(style='POP Body')
        p.add_run(text or "Não aplicável.")
        return p

    def add_bullets(items):
        items = items or []
        if not items:
            add_body_text("Não aplicável.")
            return
        for it in items:
            p = doc.add_paragraph(style='POP Bullet')
            p.add_run(f"• {it}")

    def add_std_table(headers, rows, col_widths_cm, bold_first_col=True):
        t = doc.add_table(rows=1, cols=len(headers))
        set_col_widths(t, col_widths_cm)
        hdr_cells = t.rows[0].cells
        for i, htext in enumerate(headers):
            set_cell_bg(hdr_cells[i], color_hex)
            set_cell_borders(hdr_cells[i], sz=6, color="777777")
            set_cell_margins(hdr_cells[i])
            set_cell_vcenter(hdr_cells[i])
            pp = hdr_cells[i].paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = pp.add_run(htext); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
            rr.font.size = Pt(8); rr.font.name = FONT
        trPr = t.rows[0]._tr.get_or_add_trPr()
        th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true')
        trPr.append(th)
        if not rows:
            row = t.add_row().cells
            row[0].merge(row[-1])
            set_cell_borders(row[0], sz=6, color="777777")
            set_cell_margins(row[0])
            pp = row[0].paragraphs[0]
            rr = pp.add_run("Não aplicável"); rr.font.size = Pt(8); rr.font.name = FONT
            return t
        for rdata in rows:
            row = t.add_row().cells
            for i, v in enumerate(rdata):
                set_cell_borders(row[i], sz=6, color="777777")
                set_cell_margins(row[i])
                set_cell_vcenter(row[i])
                pp = row[i].paragraphs[0]
                rr = pp.add_run(str(v)); rr.font.size = Pt(8); rr.font.name = FONT
                if bold_first_col and i == 0:
                    rr.bold = True
        return t

    usable_body_cm = usable_cm

    # 1. Objetivo
    add_heading("Objetivo")
    add_body_text(pop.get('objetivo'))

    # 2. Campo de aplicação
    add_heading("Campo de aplicação")
    add_body_text(pop.get('aplicacao'))

    # 3. Referências
    add_heading("Referências")
    add_bullets(pop.get('referencias'))

    # 4. Definições
    add_heading("Definições")
    defs = pop.get('definicoes') or []
    if not defs:
        add_body_text("Não aplicável.")
    else:
        for d in defs:
            p = doc.add_paragraph(style='POP Bullet')
            rb = p.add_run(f"• {d.get('termo','')}: "); rb.bold = True
            p.add_run(d.get('definicao', ''))

    # 5. Responsabilidades
    add_heading("Responsabilidades")
    resp = pop.get('responsaveis') or []
    resp_rows = [[r.get('papel', ''), '; '.join(r.get('tarefas') or [])] for r in resp]
    add_std_table(["Função", "Responsabilidade"], resp_rows,
                   [usable_body_cm * 0.32, usable_body_cm * 0.68])

    # 6. Materiais, utensílios, equipamentos e EPIs
    add_heading("Materiais, utensílios, equipamentos e EPIs")
    materiais = pop.get('materiais') or []
    mat_rows = []
    for m in materiais:
        if isinstance(m, dict):
            mat_rows.append([m.get('grupo', ''), m.get('item', ''), m.get('obs', '')])
        else:
            mat_rows.append(['', str(m), ''])
    add_std_table(["Grupo", "Item / quantidade", "Observação"], mat_rows,
                   [usable_body_cm * 0.20, usable_body_cm * 0.42, usable_body_cm * 0.38])

    # 7. Descrição do procedimento
    heading_n7 = add_heading("Descrição do procedimento")
    n7 = counter["n"]
    if pop.get('frequencia'):
        p = doc.add_paragraph(style='POP Body')
        rb = p.add_run("Frequência e horário: "); rb.bold = True
        p.add_run(pop['frequencia'])
    etapas = pop.get('etapas') or []
    if not etapas:
        add_body_text("Não aplicável.")
    else:
        for i, e in enumerate(etapas):
            p = doc.add_paragraph(style='POP Body')
            set_keep_with_next(p)
            destaque = e.get('destaque')
            tag_txt, tag_color = DESTAQUE_TAGS.get(destaque, ("", None))
            rb = p.add_run(f"{n7}.{i+1} {e.get('titulo','')}{tag_txt}: ")
            rb.bold = True
            if tag_color:
                rb.font.color.rgb = tag_color
            for item in (e.get('itens') or []):
                pi = doc.add_paragraph(style='POP Bullet')
                pi.add_run(f"• {item}")

    # 8. Pontos críticos de controle (PCC)
    add_heading("Pontos críticos de controle (PCC)")
    pcc = pop.get('pcc') or []
    pcc_rows = [[p.get('etapa', ''), p.get('risco', ''), p.get('controle', ''), p.get('responsavel', '')]
                for p in pcc]
    add_std_table(["Etapa", "Risco / desvio", "Medida de controle", "Responsável"], pcc_rows,
                   [usable_body_cm * 0.16, usable_body_cm * 0.28, usable_body_cm * 0.38, usable_body_cm * 0.18])

    # 9. Registros obrigatórios
    add_heading("Registros obrigatórios")
    add_bullets(pop.get('registros'))

    # 10. Indicadores, rendimento e critérios de aceitação
    add_heading("Indicadores, rendimento e critérios de aceitação")
    indicadores = pop.get('indicadores') or []
    ind_rows = []
    for it in indicadores:
        if isinstance(it, dict):
            ind_rows.append([it.get('controle', ''), it.get('criterio', '')])
        else:
            ind_rows.append(['', str(it)])
    add_std_table(["Controle", "Critério"], ind_rows,
                   [usable_body_cm * 0.28, usable_body_cm * 0.72])

    if pop.get('obs'):
        p = doc.add_paragraph(style='POP Body')
        rb = p.add_run("Observações complementares: "); rb.bold = True
        p.add_run(pop['obs'])

    # 11. Fluxo do processo — página própria, caixas + setas (padrão Word nativo)
    doc.add_page_break()
    add_heading("Fluxo do processo")
    if etapas:
        n = len(etapas)
        tint = tint_of(color_hex, amount=0.10)
        for i, e in enumerate(etapas):
            is_edge = (i == 0 or i == n - 1)
            box = doc.add_table(rows=1, cols=1)
            set_col_widths(box, [usable_body_cm])
            cell = box.rows[0].cells[0]
            set_cell_borders(cell, sz=6, color=color_hex)
            set_cell_margins(cell)
            set_cell_vcenter(cell)
            set_cell_bg(cell, color_hex if is_edge else tint)
            pp = cell.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.space_after = Pt(0)
            rr = pp.add_run(e.get('titulo') or f"Etapa {i+1}")
            rr.bold = True; rr.font.size = Pt(9); rr.font.name = FONT
            rr.font.color.rgb = RGBColor(255, 255, 255) if is_edge else RGBColor(0, 0, 0)
            if i < n - 1:
                pa = doc.add_paragraph()
                pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pa.paragraph_format.space_before = Pt(2)
                pa.paragraph_format.space_after = Pt(2)
                ra = pa.add_run("↓")
                ra.font.size = Pt(11); ra.font.color.rgb = rgb_of(color_hex); ra.font.name = FONT
    else:
        add_body_text("Não aplicável — nenhuma etapa cadastrada.")

    doc.save(out_path)
    return out_path


def docx_filename(pop):
    revisoes = pop.get('revisoes') or []
    rev_num = str(revisoes[-1].get('numero', 0)) if revisoes else '0'
    codigo = slugify(pop.get('codigo') or 'POP-0000', 20)
    titulo_curto = slugify(' '.join((pop.get('titulo') or 'POP').split()[:4]), 30)
    return f"POP_{codigo}_{titulo_curto}_REV_{rev_num}.docx"
