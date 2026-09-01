"""
RIOS — Extração de conteúdo bruto de um POP em Word (.docx) para importação via IA.

Lê parágrafos e tabelas na ordem em que aparecem no documento (cabeçalho
institucional + corpo), sem interpretar nada — só produz um texto linear que
a IA usa para mapear os campos do formulário de POP (ver IMPORT_POP_SYSTEM
em rios_server.py). A responsabilidade de "entender" o conteúdo é da IA, não
deste módulo — aqui é leitura burra e fiel do que está no arquivo.
"""
from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_block_items(parent):
    """Percorre parágrafos e tabelas na ordem original do documento (recipe
    padrão do python-docx — a API pública só expõe doc.paragraphs e doc.tables
    separados, sem preservar a ordem relativa entre eles)."""
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)


def _table_to_lines(table):
    lines = ["[TABELA]"]
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " / ") for c in row.cells]
        lines.append(" | ".join(cells))
    lines.append("[/TABELA]")
    return lines


def extract_docx_outline(path, max_chars=24000):
    """Retorna um texto único com o conteúdo do .docx em ordem de leitura:
    primeiro o cabeçalho institucional (repete em toda página — geralmente
    tem título, código, revisão, data), depois o corpo (parágrafos e tabelas
    intercalados, como aparecem no documento)."""
    doc = Document(path)
    lines = []

    try:
        header = doc.sections[0].header
        htxt = []
        for p in header.paragraphs:
            if p.text.strip():
                htxt.append(p.text.strip())
        for t in header.tables:
            htxt.extend(_table_to_lines(t))
        if htxt:
            lines.append("[CABECALHO]")
            lines.extend(htxt)
            lines.append("[/CABECALHO]")
    except Exception:
        pass

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            txt = block.text.strip()
            if txt:
                lines.append(txt)
        else:
            lines.extend(_table_to_lines(block))

    out = "\n".join(lines)
    if len(out) > max_chars:
        out = out[:max_chars] + "\n...[documento truncado — muito longo para leitura completa]"
    return out
